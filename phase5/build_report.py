"""
Generator Knowledge Discovery Report (phase5/Knowledge_Discovery_Report.docx).
Struktur mengikuti ketat `docs/question/Data Mining Report Template.docx`:
Executive Summary -> Dataset and Methodology -> Findings (3) -> Limitations -> Appendix.

Angka appendix dibaca langsung dari data/processed/ agar konsisten dengan notebook.
Jalankan:  python phase5/build_report.py
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..")
P = lambda f: os.path.join(ROOT, "data", "processed", f)

# ----------------------------------------------------------------- data
prof  = pd.read_csv(P("phase2_cluster_profiles.csv"))
rules = pd.read_csv(P("phase3_association_rules.csv"))
anom  = pd.read_csv(P("phase4_anomalies.csv"))

def clean_item(s):
    # 'int_rate_IntRate Low' -> 'int_rate = IntRate Low'
    out = []
    for it in s.split(", "):
        for pref in ("fico_range_low", "annual_inc", "loan_amnt", "revol_util",
                     "emp_length", "int_rate", "dti", "grade", "purpose", "home"):
            if it.startswith(pref + "_"):
                out.append(f"{pref} = {it[len(pref)+1:]}")
                break
        else:
            out.append(it)
    return " + ".join(out)

# ----------------------------------------------------------------- helpers
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5)

def h(text, level):
    doc.add_heading(text, level=level)

def p(text, bold=False, italic=False, size=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold; run.italic = italic
    if size: run.font.size = Pt(size)
    return par

def table(headers, rows, widths=None, font=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(htxt)); r.bold = True; r.font.size = Pt(font)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(font)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

# ================================================================= COVER
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Knowledge Discovery Report"); r.bold = True; r.font.size = Pt(22)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Segmentasi Risiko Peminjam pada Dataset Lending Club\n"
                "Penerapan proses Knowledge Discovery in Databases (KDD)")
r.font.size = Pt(12)

table(["Field", "Isi"], [
    ["Group Name",     "Kelompok 2"],
    ["Dataset",        "Lending Club Loans (accepted_2007_to_2018Q4, Kaggle)"],
    ["Domain Focus",   "Loan issuance, risk grading, borrower financial characteristics"],
    ["Submission Date","[isi tanggal submit]"],
    ["GitHub URL",     "[isi URL repo GitHub]"],
], widths=[1.6, 4.6], font=10)
doc.add_paragraph()
table(["Student Name", "Role"], [
    ["[Nama 1]", "Data Engineer"],
    ["[Nama 2]", "Data Engineer"],
    ["[Nama 3]", "Pattern Analyst"],
    ["[Nama 4]", "Segmentation Specialist"],
    ["[Nama 5]", "Insight Communicator"],
], widths=[3.1, 3.1], font=10)
doc.add_page_break()

# ================================================================= 1. EXEC SUMMARY
h("Executive Summary", 1)
p("Analisis KDD atas 2.260.668 pinjaman Lending Club menemukan bahwa risiko gagal bayar "
  "terkonsentrasi pada segmen mayoritas: 62,1 persen peminjam berada dalam satu segmen "
  "dengan tingkat charged-off 14,4 persen, hampir dua kali lipat segmen prime (7,8 persen), "
  "dan segmen ini terpisah paling tajam pada kombinasi FICO, utilisasi revolving, dan ukuran "
  "pinjaman. Dua hal yang belum tercermin dalam praktik saat ini: grade internal hampir "
  "sepenuhnya cerminan suku bunga (hanya 1 dari 1.023 aturan asosiasi yang memprediksi "
  "Grade A tanpa suku bunga), dan outlier statistik justru didominasi nasabah affluent yang "
  "membayar lancar (bad-rate 8,7 persen vs populasi 13,1 persen). Bank sebaiknya "
  "mengkalibrasi harga dan limit pada kombinasi atribut tersebut, dan mengganti flag "
  "berbasis outlier dengan aturan kombinasi atribut tekanan kredit yang tervalidasi "
  "menaikkan risiko gagal bayar 1,47 kali (19,2 persen).")

# ================================================================= 2. METHODOLOGY
h("Dataset and Methodology", 1)

h("Dataset", 2)
p("Lending Club Accepted Loans 2007–2018Q4, diunduh dari Kaggle "
  "(kaggle.com/datasets/wordsforthewise/lending-club), domain peer-to-peer lending dan "
  "risiko kredit konsumen. Seluruh 2.260.668 record dipakai (151 kolom mentah; 12 fitur "
  "final setelah feature selection). Sampling hanya dilakukan untuk komputasi yang tidak "
  "traktabel pada data penuh, selalu acak dan seeded (seed 42): mutual information 50.000 "
  "baris, UMAP 100.000, hierarchical clustering 12.000, dan Apriori 250.000.")

h("Phase 1: Data Preprocessing", 2)
p("Pembersihan tidak menghapus kolom hanya karena proporsi kosongnya tinggi. Setiap kolom "
  "bermasalah lebih dulu diberi dugaan mekanisme (structural missing, data leakage "
  "pasca-origination, atau MAR temporal) dan diverifikasi terhadap data, lalu keputusan "
  "mengikuti mekanisme itu: kolom bulan-sejak-kejadian yang kosong karena peristiwanya belum "
  "pernah terjadi diisi sentinel 999 plus flag biner, kolom leakage dibuang, kolom dengan "
  "missing di atas 40 persen dibuang, sisanya diimputasi median/modus (mort_acc diimputasi "
  "berbasis grup total_acc). Sebanyak 33 baris duplikat/tanpa label dibuang. Transformasi "
  "mencakup ordinal encoding grade dan sub_grade, binning kuantil lima fitur kontinu untuk "
  "persiapan ARM, winsorisasi data-driven pada fitur kontinu ber-|skew| di atas 1 (dengan "
  "penjaga anti-kolaps untuk fitur zero-inflated), log1p kondisional, dan standardisasi. "
  "Fitur nominal murni dibuang dari matriks clustering, bukan di-one-hot, karena one-hot "
  "mendistorsi jarak Euclidean pada clustering. "
  "Feature selection memakai dua metode: pemangkasan multikolinearitas (|r| di atas 0,85) "
  "lalu mutual information terhadap grade; hasilnya 12 fitur final. PCA sembilan komponen "
  "(90 persen varians) menjadi masukan K-Means dan hierarchical; UMAP varian densMAP menjadi "
  "masukan DBSCAN.")

h("Phase 2: Clustering", 2)
p("Tiga algoritma diterapkan: K-Means (MiniBatch) pada ruang PCA, hierarchical clustering, "
  "dan DBSCAN. Jumlah cluster ditetapkan dua setelah uji stabilitas lintas lima sample: "
  "kurva Elbow menurun mulus tanpa siku yang menonjolkan K tertentu, sedangkan Silhouette "
  "memilih K=2 secara konsisten (rata-rata 0,161, deviasi 0,010, menang di 4 dari 5 sample), "
  "sehingga keputusan bertumpu pada Silhouette. Untuk hierarchical, linkage average dan "
  "single menghasilkan chaining (satu cluster raksasa plus singleton) sehingga linkage Ward "
  "dipilih (cophenetic correlation 0,360; cophenetic average/single lebih tinggi tetapi "
  "strukturnya degenerate). DBSCAN pada embedding densMAP memakai eps 0,471 dan min_samples "
  "12, keduanya data-driven (knee k-distance; pembulatan ln n), menghasilkan 501 noise "
  "point (0,50 persen). Adjusted Rand Index K-Means terhadap hierarchical 0,357 (moderat).")

h("Phase 3: Association Rule Mining", 2)
p("Variabel kontinu didiskretisasi memakai ambang berbasis domain, bukan kuantil sembarang: "
  "bracket resmi FICO, ambang DTI 43 persen dari aturan Qualified Mortgage, batas utilisasi "
  "revolving 30 persen rekomendasi biro kredit, serta tier pendapatan dan jumlah pinjaman "
  "kredit konsumen AS. Mining dijalankan pada sample acak seeded 250.000 baris (proporsi "
  "seluruh tahun origination terjaga). Apriori dengan minimum support 0,04 dan panjang "
  "itemset maksimum empat menghasilkan 1.947 frequent itemset, yang menurunkan 14.764 rule "
  "sebelum filtering. Setelah filter lift di atas 1,15 dan confidence di atas 0,5, tersisa "
  "1.023 rule (6,9 persen) dengan rentang lift 1,15 sampai 5,35.")

h("Phase 4: Anomaly Detection", 2)
p("Deteksi mencari tiga tipe outlier: point (ekstrem satu dimensi), contextual (janggal "
  "hanya secara multivariat), dan collective (region kepadatan rendah). Empat metode dengan "
  "ambang eksplisit: IQR pengali 3, Z-score ambang 3, jarak Mahalanobis pada persentil "
  "chi-square 99,9 persen, dan Isolation Forest dengan contamination 1 persen. Deteksi "
  "berjalan pada data non-winsorized agar nilai asli tidak terhapus. Anomali terkonfirmasi "
  "bila ditandai minimal dua metode: dari 271.143 kandidat, 116.225 terkonfirmasi, lalu "
  "disilangkan dengan noise DBSCAN Phase 2.")

# ================================================================= 3. FINDINGS
h("Findings", 1)
p("Setiap temuan berasal dari fase berbeda, memuat angka dari analisis, tidak terlihat "
  "langsung dari data mentah, mengimplikasikan tindakan bank spesifik, dan menghindari "
  "klaim kausal.", italic=True)

h("Finding 1 — Risiko gagal bayar terkonsentrasi pada segmen mayoritas (Clustering)", 2)
p("62,1 persen peminjam berada di satu segmen dengan tingkat charged-off 14,4 persen, "
  "hampir dua kali lipat segmen prime (7,8 persen).", bold=True)
h("Evidence", 3)
p("Segmentasi K-Means membelah populasi menjadi dua. Cluster 0 (dinamai Higher-risk) "
  "mencakup 62,1 persen peminjam dengan tingkat charged-off 14,4 persen; Cluster 1 (Prime) "
  "mencakup 37,9 persen dengan charged-off 7,8 persen. Tingkat di sini dihitung dari status "
  "Charged Off saja (rata-rata populasi 11,9 persen). Dalam satuan simpangan baku (selisih "
  "z-score antar cluster, konsisten dengan radar profil Phase 2), pembeda paling tajam "
  "adalah FICO (1,15 SD; 684 berbanding 722), disusul ukuran pinjaman (0,87 SD; 12 ribu "
  "berbanding 20 ribu dolar) dan utilisasi revolving (0,85 SD; 58 berbanding 37 persen); "
  "gap penghasilan besar secara nominal (63 ribu berbanding 102 ribu dolar) tetapi hanya "
  "0,35 SD karena varians penghasilan sangat lebar.")
h("Corroboration", 3)
p("Hierarchical clustering linkage Ward menghasilkan pembagian dua kelompok serupa "
  "(Adjusted Rand Index 0,357, moderat); tabel silang menempatkan mayoritas anggota pada "
  "sel yang bersesuaian. Silhouette 0,161 memang rendah, wajar untuk data sosial-ekonomi "
  "kontinu tanpa gap tajam, tetapi kestabilan pemisahan lintas lima sample memperkuat bahwa "
  "dua segmen ini nyata.")
h("Business Implication", 3)
p("Intuisi umum menduga peminjam bermasalah adalah minoritas kecil. Data menunjukkan "
  "sebaliknya: hampir dua pertiga buku pinjaman berada di segmen yang lebih berisiko, dan "
  "risiko itu terbaca dari beberapa atribut sekaligus, bukan satu skor tunggal: utilisasi "
  "revolving dan ukuran pinjaman menambah pemisahan hampir sebesar FICO.")
h("Recommended Action", 3)
p("Tim kebijakan kredit sebaiknya mengkalibrasi harga dan limit pada kombinasi FICO, "
  "utilisasi revolving, dan ukuran pinjaman, bukan FICO tunggal. Karena segmen Higher-risk "
  "berukuran besar, perbaikan kecil pada penetapan harga segmen ini berdampak jauh lebih "
  "besar pada kerugian portofolio dibanding menyaring minoritas ekor.")

h("Finding 2 — Grade internal hampir sepenuhnya cerminan suku bunga (ARM)", 2)
p("Hanya 1 dari 1.023 rule yang memprediksi Grade A tanpa suku bunga, dan satu-satunya "
  "prediktornya adalah FICO 740–799 (confidence 57 persen, lift 2,99).", bold=True)
h("Evidence", 3)
p("Dari 1.023 rule retained, seluruh rule kuat menuju grade memuat kategori suku bunga; "
  "satu-satunya pengecualian adalah FICO Very Good (740–799) menuju Grade A dengan "
  "confidence 0,57 dan lift 2,99, dan FICO memang input langsung model grading. Rule "
  "berlift tertinggi (sampai 5,35) semuanya memetakan pasangan grade dan suku bunga, "
  "misalnya Grade A menuju suku bunga rendah pada confidence 1,00 (lift 4,00). Kombinasi "
  "atribut fundamental (DTI sehat, FICO baik, utilisasi rendah) berasosiasi dengan Grade A "
  "pada 56 rule mining-angle, tetapi selalu bersama kategori suku bunga.")
h("Corroboration", 3)
p("Pola grade-suku bunga konsisten monoton di seluruh tingkat: Grade B dengan bunga medium "
  "(lift 2,92), Grade C dengan bunga tinggi (3,03), Grade D/E dengan bunga sangat tinggi "
  "(3,68–4,06), dengan confidence 0,73 sampai 1,00. Contoh rule pada dokumen soal yang "
  "melibatkan small_business tidak didukung data (0 rule), sehingga interpretasi dibangun "
  "dari rule aktual, bukan asumsi.")
h("Business Implication", 3)
p("Grade Lending Club pada praktiknya adalah fungsi dari penetapan harga internal (dan "
  "FICO), bukan ringkasan independen kelayakan kredit. Memakai grade dan suku bunga "
  "bersamaan dalam model hilir berarti memasukkan informasi yang sama dua kali.")
h("Recommended Action", 3)
p("Tim pemodelan risiko sebaiknya tidak memperlakukan grade sebagai fitur independen di "
  "samping suku bunga; modelkan langsung dari atribut fundamental peminjam (DTI, FICO, "
  "utilisasi, penghasilan) agar penilaian tidak mengulang keputusan harga yang sudah dibuat.")

h("Finding 3 — Anomali statistik bukan penanda risiko, kecuali satu sub-kelompok (Anomaly Detection)", 2)
p("Anomali terkonfirmasi gagal bayar lebih jarang dari populasi (8,7 vs 13,1 persen), "
  "tetapi sub-kelompok high-stress berjumlah 11.358 baris gagal bayar 19,2 persen "
  "(1,47 kali populasi).", bold=True)
h("Evidence", 3)
p("Bad-rate di temuan ini memakai definisi luas (Charged Off, Default, dan Late; populasi "
  "13,1 persen), berbeda dari Finding 1 yang memakai Charged Off saja. Anomali "
  "terkonfirmasi justru gagal bayar lebih jarang (8,7 persen) karena mayoritas outlier "
  "adalah peminjam berpenghasilan sangat tinggi atau berlimit kartu besar yang membayar "
  "lancar. Sub-kelompok high-stress dibentuk independen dari skor gabungan DTI, utilisasi "
  "revolving, utilisasi total, dan jumlah inquiry (tanpa menyentuh status pinjaman): "
  "11.358 baris ini gagal bayar 19,2 persen.")
h("Corroboration", 3)
p("Karena sub-kelompok dibentuk sebelum melihat outcome, angka 19,2 persen adalah validasi, "
  "bukan definisi berputar. Deteksi saling menguatkan lintas metode: dari 501 noise DBSCAN "
  "Phase 2, 66 persen ikut ditandai metode statistik atau Isolation Forest. Terpisah dari "
  "itu, 2.654 anomali terklasifikasi data error murni (misalnya penghasilan 110 juta dolar, "
  "atau DTI bernilai sentinel 999).")
h("Business Implication", 3)
p("Menyaring nasabah hanya karena ia outlier statistik salah sasaran: kelompok itu justru "
  "berisi nasabah bernilai tinggi. Sinyal risiko sesungguhnya datang dari kombinasi atribut "
  "tekanan kredit, bukan dari status keanehan itu sendiri.")
h("Recommended Action", 3)
p("Tim underwriting sebaiknya memakai kombinasi atribut high-stress (DTI, utilisasi, dan "
  "inquiry tinggi bersamaan) sebagai aturan flag untuk review manual, bukan menandai baris "
  "karena outlier. Baris data error dikoreksi atau dikeluarkan lebih dulu agar tidak "
  "mendistorsi model.")

# ================================================================= 4. LIMITATIONS
h("Limitations", 1)
h("Scope of outlier detection", 2)
p("Deteksi anomali berjalan pada 12 fitur numerik terpilih dan sebagian besar menemukan "
  "point outlier (115.945 dari 116.225); contextual murni (43) dan collective (237) sangat "
  "sedikit, sebagian karena fitur finansial berekor berat pada dimensi tunggal, sebagian "
  "karena keterbatasan ruang fitur. Anomali yang hanya muncul pada fitur yang dibuang saat "
  "feature selection tidak tertangkap.")
h("Correlation versus causation", 2)
p("Semua temuan bersifat asosiatif. Cluster Higher-risk berasosiasi dengan penghasilan "
  "rendah dan utilisasi tinggi, tetapi analisis ini tidak membuktikan sebab-akibat; kami "
  "melaporkan selisih tingkat, bukan efek kausal.")
h("Dataset representativeness", 2)
p("Data hanya memuat pinjaman yang disetujui Lending Club 2007–2018 sehingga mengandung "
  "bias seleksi (pemohon ditolak tidak terlihat), mencakup kondisi ekonomi periode itu "
  "saja, dan grade/suku bunga mengikuti kebijakan internal Lending Club. Komputasi berat "
  "(MI, UMAP, hierarchical, Apriori) berjalan pada sample acak seeded, bukan data penuh.")
h("What additional data would improve these findings", 2)
p("Data pemohon yang ditolak akan memperbaiki analisis seleksi; riwayat pembayaran bulanan "
  "memungkinkan analisis temporal yang lebih dalam daripada status akhir; indikator "
  "makroekonomi per periode origination membantu memisahkan risiko peminjam dari risiko "
  "siklus; dan atribut biro kredit yang lebih rinci memungkinkan penilaian risiko yang "
  "independen dari penetapan harga internal.")

# ================================================================= 5. APPENDIX
doc.add_page_break()
h("Appendix", 1)

h("Appendix A — Full Cluster Profiles", 2)
rows = []
for _, r0 in prof.iterrows():
    rows.append([int(r0["cluster"]), r0["profil"], r0["size_%"], r0["CO_%"],
                 f'{r0["annual_inc"]:,.0f}', f'{r0["fico_range_low"]:.0f}',
                 r0["revol_util"], r0["dti"], f'{r0["loan_amnt"]:,.0f}',
                 r0["grade"], r0["term"], r0["emp_length"]])
table(["Cluster", "Profil", "Ukuran %", "CO %", "Income", "FICO", "Revol util",
       "DTI", "Loan amnt", "Grade", "Term", "Emp len"], rows, font=8)
p("CO % = tingkat Charged Off (definisi sempit). Profil outlier DBSCAN (501 noise): "
  "income rata-rata 117.758 dolar, DTI 25,7, loan 21.459 dolar, charged-off 15,8 persen "
  "(populasi 11,9) — profil 'stretched high-earners'.", size=9)

h("Appendix B — Association Rules (40 lift tertinggi)", 2)
p("Tabel lengkap 1.023 rule retained tersedia di data/processed/phase3_association_rules.csv, "
  "terurut menurun berdasarkan lift. Berikut 40 teratas.", size=9)
top40 = rules.nlargest(40, "lift")
rows = [[clean_item(r0["antecedents"]), clean_item(r0["consequents"]),
         f'{r0["support"]:.3f}', f'{r0["confidence"]:.2f}', f'{r0["lift"]:.2f}']
        for _, r0 in top40.iterrows()]
table(["Antecedent", "Consequent", "Supp", "Conf", "Lift"], rows,
      widths=[2.9, 2.0, 0.55, 0.55, 0.55], font=7)

h("Appendix C — Anomaly Detection Results (sampel per kategori)", 2)
p("Seluruh 116.225 anomali terkonfirmasi tersimpan di data/processed/phase4_anomalies.csv "
  "(kolom: flag per metode, skor Mahalanobis, risk score, tipe outlier, klasifikasi). "
  "Berikut record teratas per kategori menurut jarak Mahalanobis, dengan interpretasi "
  "bisnis per record.", size=9)
interp = {
    "data_error": "Income tidak plausibel utk pinjaman konsumen; salah entri — koreksi/keluarkan",
    "risk_signal": "Tekanan kredit tinggi (DTI/utilisasi/inquiry) — flag review manual underwriting",
    "rare_case": "Affluent/limit besar, plausibel & mayoritas lancar — jangan tolak otomatis",
}
show = ["orig_index", "annual_inc", "dti", "fico_range_low", "revol_util",
        "loan_amnt", "loan_status", "n_methods", "outlier_type"]
rows = []
for cat in ["data_error", "risk_signal", "rare_case"]:
    sub = anom[anom["classification"] == cat].sort_values("maha_dist2", ascending=False).head(6)
    for _, r0 in sub.iterrows():
        rows.append([cat, int(r0["orig_index"]), f'{r0["annual_inc"]:,.0f}', r0["dti"],
                     f'{r0["fico_range_low"]:.0f}', r0["revol_util"], f'{r0["loan_amnt"]:,.0f}',
                     str(r0["loan_status"])[:28], int(r0["n_methods"]), r0["outlier_type"],
                     interp[cat]])
table(["Kategori", "ID", "Income", "DTI", "FICO", "Util", "Loan", "Status", "Met.",
       "Tipe", "Interpretasi bisnis"], rows,
      widths=[0.7, 0.55, 0.75, 0.4, 0.4, 0.4, 0.55, 1.0, 0.35, 0.6, 1.6], font=6.5)

h("Appendix D — Evaluation Metrics Summary", 2)
table(["Metrik", "Nilai"], [
    ["Silhouette Score (K-Means, K final = 2)", "0,161 (± 0,010 lintas 5 sample)"],
    ["Cophenetic Correlation (linkage Ward)", "0,360"],
    ["Adjusted Rand Index (K-Means vs Hierarchical)", "0,357"],
    ["Frequent itemset (Apriori, support ≥ 0,04)", "1.947"],
    ["Rule dihasilkan sebelum filtering", "14.764"],
    ["Rule dipertahankan setelah filter (lift > 1,15; conf > 0,5)", "1.023"],
    ["Lift tertinggi pada retained rules", "5,35"],
    ["Total kandidat anomali sebelum korroborasi (≥1 metode)", "271.143"],
    ["Anomali terkonfirmasi (≥2 metode)", "116.225"],
    ["Precision/Recall vs label", "Tidak berlaku (unsupervised; label outcome hanya untuk validasi)"],
], widths=[4.0, 2.6], font=9)

out = os.path.join(ROOT, "phase5", "Knowledge_Discovery_Report.docx")
doc.save(out)
print("Tersimpan:", out)
